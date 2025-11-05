#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GUI Question Converter - Giao diện chuyển đổi câu hỏi
Panel trái: Bảng câu hỏi và đáp án
Panel phải: Nội dung file gốc với highlight
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document
import re
import os
import xml.etree.ElementTree as ET

class QuestionConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Question Converter - Chuyển đổi câu hỏi")
        self.root.geometry("1400x800")
        self.root.configure(bg='#f0f0f0')
        
        # Dữ liệu
        self.questions = []
        self.raw_content = []
        self.current_file = None
        
        # Thiết lập giao diện
        self.setup_ui()
        
        # Load dữ liệu mẫu
        self.load_sample_data()
    
    def setup_ui(self):
        """Thiết lập giao diện người dùng"""
        # Menu bar
        self.create_menu()
        
        # Toolbar
        self.create_toolbar()
        
        # Main panels
        self.create_main_panels()
        
        # Status bar
        self.create_status_bar()
    
    def create_menu(self):
        """Tạo menu bar"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Mở file (.doc/.docx/.txt/.xml)", command=self.open_file)
        file_menu.add_separator()
        file_menu.add_command(label="Xuất ra file Word", command=self.export_to_word)
        file_menu.add_separator()
        file_menu.add_command(label="Thoát", command=self.root.quit)
        
        # View menu
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="View", menu=view_menu)
        view_menu.add_command(label="Refresh", command=self.refresh_view)
        view_menu.add_command(label="Kiểm tra dữ liệu", command=self.check_data_quality)
        view_menu.add_separator()
        view_menu.add_command(label="🔧 Sửa số thứ tự (226a→227)", command=self.fix_question_numbers)
        
        # Export menu
        export_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Export", menu=export_menu)
        export_menu.add_command(label="Xuất ra Word (.docx)", command=self.export_to_word)
        export_menu.add_command(label="Xuất ra JSON", command=self.export_to_json)
        export_menu.add_command(label="Xuất ra TXT", command=self.export_to_txt)
        export_menu.add_command(label="Xuất ra XML", command=self.export_to_xml)
    
    def create_toolbar(self):
        """Tạo toolbar"""
        toolbar_frame = tk.Frame(self.root, bg='#e0e0e0', height=40)
        toolbar_frame.pack(fill=tk.X, padx=5, pady=2)
        toolbar_frame.pack_propagate(False)
        
        # Buttons
        tk.Button(toolbar_frame, text="📁 Mở File", command=self.open_file, 
                 bg='#4CAF50', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="🔄 Refresh", command=self.refresh_view,
                 bg='#2196F3', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="� Kiểm tra", command=self.check_data_quality,
                 bg='#E91E63', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="�💾 Xuất Word", command=self.export_to_word,
                 bg='#FF9800', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="📄 Xuất JSON", command=self.export_to_json,
                 bg='#9C27B0', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="📝 Xuất TXT", command=self.export_to_txt,
                 bg='#607D8B', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        tk.Button(toolbar_frame, text="📋 Xuất XML", command=self.export_to_xml,
                 bg='#009688', fg='white', padx=10).pack(side=tk.LEFT, padx=5, pady=5)
        
        # Info label
        self.info_label = tk.Label(toolbar_frame, text="Sẵn sàng - Dữ liệu mẫu đã được tải", 
                                  bg='#e0e0e0', fg='#333')
        self.info_label.pack(side=tk.RIGHT, padx=10, pady=5)
    
    def create_main_panels(self):
        """Tạo notebook với các tab"""
        # Main container
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Tạo Notebook (tabs)
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Tab 1: Converter (giao diện cũ)
        self.create_converter_tab()
        
        # Tab 2: Trộn đề
        self.create_exam_mixer_tab()
    
    def create_converter_tab(self):
        """Tab chuyển đổi câu hỏi (giao diện cũ)"""
        converter_frame = tk.Frame(self.notebook)
        self.notebook.add(converter_frame, text="📝 Chuyển đổi câu hỏi")
        
        # Left panel - Question table
        left_frame = tk.LabelFrame(converter_frame, text="📋 Danh sách câu hỏi", 
                                  font=('Arial', 10, 'bold'), fg='#2196F3')
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        self.create_question_table(left_frame)
        
        # Right panel - Original content
        right_frame = tk.LabelFrame(converter_frame, text="📄 Nội dung file gốc", 
                                   font=('Arial', 10, 'bold'), fg='#FF9800')
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        self.create_content_viewer(right_frame)
    
    def create_question_table(self, parent):
        """Tạo bảng câu hỏi ở panel trái với format 3 cột: Label, Content, Group"""
        # Treeview để hiển thị bảng với 3 cột
        columns = ('Label', 'Content', 'Group')
        self.tree = ttk.Treeview(parent, columns=columns, show='headings', height=15)
        
        # Định nghĩa headers
        self.tree.heading('Label', text='')
        self.tree.heading('Content', text='Nội dung')
        self.tree.heading('Group', text='Nhóm')
        
        # Độ rộng cột
        self.tree.column('Label', width=100, anchor='w')
        self.tree.column('Content', width=500, anchor='w')
        self.tree.column('Group', width=200, anchor='w')
        
        # Scrollbars cho bảng
        tree_scroll_y = ttk.Scrollbar(parent, orient="vertical", command=self.tree.yview)
        tree_scroll_x = ttk.Scrollbar(parent, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        # Pack
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree_scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        tree_scroll_x.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Bind sự kiện click
        self.tree.bind('<<TreeviewSelect>>', self.on_question_select)
    
    def create_content_viewer(self, parent):
        """Tạo viewer hiển thị nội dung file gốc ở panel phải"""
        # Frame chứa line numbers và content
        viewer_frame = tk.Frame(parent)
        viewer_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Line number widget (không có scrollbar)
        self.line_numbers = tk.Text(viewer_frame, 
                                    width=5, 
                                    padx=3,
                                    takefocus=0,
                                    border=0,
                                    background='#f0f0f0',
                                    foreground='#666',
                                    state='disabled',
                                    font=('Consolas', 11))
        self.line_numbers.pack(side=tk.LEFT, fill=tk.Y)
        
        # Text widget để hiển thị nội dung (có scrollbar)
        self.content_text = scrolledtext.ScrolledText(viewer_frame, 
                                                     wrap=tk.WORD, 
                                                     width=50, 
                                                     height=30,
                                                     font=('Consolas', 11))
        self.content_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Đồng bộ scroll giữa line numbers và content
        self.content_text.config(yscrollcommand=self.on_text_scroll)
        
        # Configure tags cho highlight
        self.content_text.tag_configure('highlight', background='yellow', foreground='black', font=('Consolas', 11, 'bold'))
        self.content_text.tag_configure('question', background='lightblue', foreground='black', font=('Consolas', 11, 'bold'))
        self.content_text.tag_configure('correct_answer', background='lightgreen', foreground='black', font=('Consolas', 11, 'bold'))
        self.content_text.tag_configure('error', background='red', foreground='white', font=('Consolas', 11, 'bold'))
        self.content_text.tag_configure('error_text', background='#ffcccc', foreground='darkred', font=('Consolas', 11))
        self.content_text.tag_configure('error_highlight', background='#ff6666', foreground='white', font=('Consolas', 11, 'bold'))
        
        # Thêm context menu cho copy/paste
        self.create_context_menu()
        
    def create_context_menu(self):
        """Tạo context menu cho content_text"""
        self.context_menu = tk.Menu(self.content_text, tearoff=0)
        self.context_menu.add_command(label="Copy (Ctrl+C)", command=self.copy_text)
        self.context_menu.add_command(label="Paste (Ctrl+V)", command=self.paste_text)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Select All (Ctrl+A)", command=self.select_all_text)
        
        # Bind right click
        self.content_text.bind("<Button-3>", self.show_context_menu)
        
        # Bind keyboard shortcuts
        self.content_text.bind("<Control-c>", lambda e: self.copy_text())
        self.content_text.bind("<Control-v>", lambda e: self.paste_text())
        self.content_text.bind("<Control-a>", lambda e: self.select_all_text())
    
    def show_context_menu(self, event):
        """Hiển thị context menu"""
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()
    
    def copy_text(self):
        """Copy text đã chọn"""
        try:
            selected_text = self.content_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)
        except tk.TclError:
            pass  # Không có text được chọn
    
    def paste_text(self):
        """Paste text từ clipboard"""
        try:
            clipboard_text = self.root.clipboard_get()
            # Chèn tại vị trí con trỏ hoặc thay thế text đã chọn
            try:
                self.content_text.delete(tk.SEL_FIRST, tk.SEL_LAST)
            except tk.TclError:
                pass  # Không có selection
            self.content_text.insert(tk.INSERT, clipboard_text)
        except tk.TclError:
            pass  # Clipboard trống
    
    def select_all_text(self):
        """Chọn toàn bộ text"""
        self.content_text.tag_add(tk.SEL, "1.0", tk.END)
        self.content_text.mark_set(tk.INSERT, "1.0")
        self.content_text.see(tk.INSERT)
        return 'break'  # Ngăn event tiếp tục
    
    def create_exam_mixer_tab(self):
        """Tab trộn đề thi"""
        mixer_frame = tk.Frame(self.notebook, bg='#f5f5f5')
        self.notebook.add(mixer_frame, text="🎲 Trộn đề thi")
        
        # Container chính
        main_container = tk.Frame(mixer_frame, bg='#f5f5f5')
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Title
        title_label = tk.Label(main_container, 
                              text="🎲 CÔNG CỤ TRỘN ĐỀ THI",
                              font=('Arial', 16, 'bold'),
                              bg='#f5f5f5',
                              fg='#1976D2')
        title_label.pack(pady=(0, 20))
        
        # Frame cho settings
        settings_frame = tk.LabelFrame(main_container, 
                                      text="⚙️ Cài đặt",
                                      font=('Arial', 11, 'bold'),
                                      bg='#ffffff',
                                      fg='#333',
                                      padx=20,
                                      pady=15)
        settings_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Số đề
        row1 = tk.Frame(settings_frame, bg='#ffffff')
        row1.pack(fill=tk.X, pady=5)
        tk.Label(row1, text="Số đề cần tạo:", font=('Arial', 10), bg='#ffffff', width=20, anchor='w').pack(side=tk.LEFT)
        self.num_exams_var = tk.StringVar(value="5")
        tk.Entry(row1, textvariable=self.num_exams_var, font=('Arial', 10), width=15).pack(side=tk.LEFT, padx=10)
        tk.Label(row1, text="(VD: 5, 10, 20)", font=('Arial', 9), fg='#666', bg='#ffffff').pack(side=tk.LEFT)
        
        # Số câu mỗi đề
        row2 = tk.Frame(settings_frame, bg='#ffffff')
        row2.pack(fill=tk.X, pady=5)
        tk.Label(row2, text="Số câu mỗi đề:", font=('Arial', 10), bg='#ffffff', width=20, anchor='w').pack(side=tk.LEFT)
        self.num_questions_var = tk.StringVar(value="30")
        tk.Entry(row2, textvariable=self.num_questions_var, font=('Arial', 10), width=15).pack(side=tk.LEFT, padx=10)
        tk.Label(row2, text="(VD: 20, 30, 50)", font=('Arial', 9), fg='#666', bg='#ffffff').pack(side=tk.LEFT)
        
        # Thông tin về dữ liệu hiện tại
        info_frame = tk.LabelFrame(main_container,
                                  text="ℹ️ Thông tin dữ liệu",
                                  font=('Arial', 11, 'bold'),
                                  bg='#E3F2FD',
                                  fg='#333',
                                  padx=20,
                                  pady=15)
        info_frame.pack(fill=tk.X, pady=(0, 15))
        
        self.mixer_info_label = tk.Label(info_frame,
                                        text="Chưa có dữ liệu. Vui lòng load file ở tab 'Chuyển đổi câu hỏi'",
                                        font=('Arial', 10),
                                        bg='#E3F2FD',
                                        fg='#555',
                                        justify=tk.LEFT)
        self.mixer_info_label.pack(anchor='w')
        
        # Buttons
        button_frame = tk.Frame(main_container, bg='#f5f5f5')
        button_frame.pack(pady=20)
        
        tk.Button(button_frame,
                 text="🎲 Tạo đề thi",
                 command=self.generate_exams,
                 font=('Arial', 11, 'bold'),
                 bg='#4CAF50',
                 fg='white',
                 padx=30,
                 pady=10,
                 cursor='hand2').pack(side=tk.LEFT, padx=10)
        
        tk.Button(button_frame,
                 text="🔄 Làm mới thông tin",
                 command=self.update_mixer_info,
                 font=('Arial', 11),
                 bg='#2196F3',
                 fg='white',
                 padx=20,
                 pady=10,
                 cursor='hand2').pack(side=tk.LEFT, padx=10)
        
        # Preview area
        preview_frame = tk.LabelFrame(main_container,
                                     text="📋 Kết quả",
                                     font=('Arial', 11, 'bold'),
                                     bg='#ffffff',
                                     fg='#333',
                                     padx=10,
                                     pady=10)
        preview_frame.pack(fill=tk.BOTH, expand=True)
        
        self.mixer_result_text = scrolledtext.ScrolledText(preview_frame,
                                                          font=('Consolas', 9),
                                                          wrap=tk.WORD,
                                                          bg='#fafafa')
        self.mixer_result_text.pack(fill=tk.BOTH, expand=True)
        
        # Load thông tin ban đầu
        self.update_mixer_info()
    
    def on_text_scroll(self, *args):
        """Callback khi content_text scroll - đồng bộ với line numbers"""
        # Cập nhật scrollbar
        if hasattr(self.content_text, 'vbar'):
            self.content_text.vbar.set(*args)
        
        # Đồng bộ scroll của line numbers
        first_visible = self.content_text.index("@0,0")
        self.line_numbers.yview_moveto(args[0])
    
    def create_status_bar(self):
        """Tạo status bar"""
        status_frame = tk.Frame(self.root, bg='#333', height=25)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(status_frame, text="Sẵn sàng", 
                                   bg='#333', fg='white', anchor='w')
        self.status_label.pack(fill=tk.X, padx=10)
    
    def update_mixer_info(self):
        """Cập nhật thông tin về dữ liệu hiện có"""
        if not self.questions:
            info_text = "❌ Chưa có dữ liệu.\n\nVui lòng load file ở tab 'Chuyển đổi câu hỏi' trước."
            self.mixer_info_label.config(text=info_text, fg='#d32f2f')
            return
        
        # Thống kê groups
        groups = {}
        for q in self.questions:
            group = q.get('group', 'Không có nhóm')
            if group not in groups:
                groups[group] = 0
            groups[group] += 1
        
        total = len(self.questions)
        num_groups = len(groups)
        
        info_text = f"✅ Tổng số câu hỏi: {total}\n"
        info_text += f"📊 Số nhóm: {num_groups}\n\n"
        info_text += "Chi tiết từng nhóm:\n"
        
        for idx, (group_name, count) in enumerate(groups.items(), 1):
            info_text += f"  {idx}. {group_name}: {count} câu\n"
        
        self.mixer_info_label.config(text=info_text, fg='#1976D2')
    
    def generate_exams(self):
        """Tạo các đề thi ngẫu nhiên"""
        import random
        
        # Validate input
        if not self.questions:
            messagebox.showerror("Lỗi", "Chưa có dữ liệu câu hỏi!\n\nVui lòng load file ở tab 'Chuyển đổi câu hỏi' trước.")
            return
        
        try:
            num_exams = int(self.num_exams_var.get())
            num_questions = int(self.num_questions_var.get())
        except ValueError:
            messagebox.showerror("Lỗi", "Vui lòng nhập số hợp lệ!")
            return
        
        if num_exams <= 0 or num_questions <= 0:
            messagebox.showerror("Lỗi", "Số đề và số câu phải lớn hơn 0!")
            return
        
        total_available = len(self.questions)
        
        # Kiểm tra đủ câu không
        if num_questions > total_available:
            messagebox.showerror("Lỗi", 
                               f"Không đủ câu hỏi!\n\n"
                               f"Số câu có sẵn: {total_available}\n"
                               f"Số câu yêu cầu: {num_questions}")
            return
        
        # Nhóm câu hỏi theo group
        groups = {}
        for q in self.questions:
            group = q.get('group', 'Không có nhóm')
            if group not in groups:
                groups[group] = []
            groups[group].append(q)
        
        num_groups = len(groups)
        
        # Tạo thư mục output
        output_folder = filedialog.askdirectory(title="Chọn thư mục lưu các đề thi")
        if not output_folder:
            return
        
        # Tạo các đề
        self.mixer_result_text.delete(1.0, tk.END)
        self.mixer_result_text.insert(tk.END, f"Đang tạo {num_exams} đề thi, mỗi đề {num_questions} câu...\n\n")
        self.root.update()
        
        created_exams = []
        
        for exam_num in range(1, num_exams + 1):
            # Tạo pool câu hỏi cho đề này
            selected_questions = self.select_questions_for_exam(groups, num_questions, num_groups)
            
            if not selected_questions:
                messagebox.showerror("Lỗi", f"Không thể tạo đề số {exam_num}!")
                break
            
            # Shuffle câu hỏi
            random.shuffle(selected_questions)
            
            # Tạo file Word
            file_name = f"De_thi_{exam_num:02d}.docx"
            file_path = os.path.join(output_folder, file_name)
            
            self.create_exam_word_file(file_path, selected_questions, exam_num)
            created_exams.append(file_name)
            
            # Cập nhật progress
            self.mixer_result_text.insert(tk.END, f"✅ Đã tạo: {file_name}\n")
            self.root.update()
        
        # Thông báo hoàn thành
        self.mixer_result_text.insert(tk.END, f"\n{'='*50}\n")
        self.mixer_result_text.insert(tk.END, f"🎉 HOÀN THÀNH!\n")
        self.mixer_result_text.insert(tk.END, f"Đã tạo {len(created_exams)} đề thi tại:\n{output_folder}\n")
        
        messagebox.showinfo("Thành công", 
                          f"Đã tạo {len(created_exams)} đề thi!\n\n"
                          f"Vị trí: {output_folder}")
    
    def select_questions_for_exam(self, groups, num_questions, num_groups):
        """Chọn câu hỏi cho một đề thi - phân đều từ các groups"""
        import random
        
        selected = []
        group_names = list(groups.keys())
        
        if num_questions <= num_groups:
            # Ít câu hơn số groups: chọn ngẫu nhiên một số groups
            selected_groups = random.sample(group_names, num_questions)
            for group in selected_groups:
                if groups[group]:
                    selected.append(random.choice(groups[group]))
        else:
            # Nhiều câu hơn số groups: phân đều
            # Bước 1: Chọn ít nhất 1 câu từ mỗi group
            for group in group_names:
                if groups[group]:
                    selected.append(random.choice(groups[group]))
            
            # Bước 2: Phân đều số câu còn lại
            remaining = num_questions - len(selected)
            questions_per_group = remaining // num_groups
            extra_questions = remaining % num_groups
            
            # Tạo pool các câu chưa được chọn
            available_by_group = {}
            for group, questions in groups.items():
                # Lọc ra các câu chưa được chọn
                available = [q for q in questions if q not in selected]
                if available:
                    available_by_group[group] = available
            
            # Phân đều câu hỏi
            for group in group_names:
                if group not in available_by_group:
                    continue
                
                # Số câu cần lấy từ group này
                num_to_take = questions_per_group
                if extra_questions > 0:
                    num_to_take += 1
                    extra_questions -= 1
                
                # Lấy câu (không vượt quá số câu có sẵn)
                num_to_take = min(num_to_take, len(available_by_group[group]))
                selected.extend(random.sample(available_by_group[group], num_to_take))
                
                if len(selected) >= num_questions:
                    break
        
        return selected[:num_questions]
    
    def create_exam_word_file(self, file_path, questions, exam_number):
        """Tạo file Word cho một đề thi với format yêu cầu"""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Tiêu đề
        heading = doc.add_heading(f'ĐỀ THI SỐ {exam_number:02d}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm câu hỏi
        for idx, q in enumerate(questions, 1):
            # Câu hỏi
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(f"Câu {idx}. {q['question']}")
            question_run.bold = True
            
            # Đáp án
            for letter in sorted(q['answers'].keys()):
                answer_para = doc.add_paragraph(f"   {letter}. {q['answers'][letter]}")
            
            # Thêm khoảng cách
            doc.add_paragraph()
        
        # Ngắt trang trước bảng đáp án
        doc.add_page_break()
        
        # Tiêu đề bảng đáp án
        answer_heading = doc.add_heading('ĐÁP ÁN', level=1)
        answer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tạo bảng đáp án (2 cột)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Câu hỏi'
        header_cells[1].text = 'Đáp án'
        
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        
        # Thêm đáp án vào bảng
        for idx, q in enumerate(questions, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = f"Câu {idx}"
            row_cells[1].text = q.get('correct_answer', 'Chưa xác định')
            
            # Center alignment
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Điều chỉnh độ rộng cột
        for i, width in enumerate([Inches(2), Inches(2)]):
            for row in table.rows:
                row.cells[i].width = width
        
        doc.save(file_path)
    
    def load_sample_data(self):
        """Load dữ liệu mẫu"""
        self.raw_content = [
            "<Gr> Kiến thức Toán học cơ bản",
            "<CH>",
            "Câu 1. Kết quả của phép tính 15 + 27 là:",
            "A. 40",
            "*B. 42",
            "C. 44",
            "D. 45",
            "</CH>",
            "<CH>",
            "Câu 2. Diện tích hình chữ nhật có chiều dài 8cm, chiều rộng 5cm là:",
            "A. 13 cm²",
            "B. 26 cm²",
            "*C. 40 cm²",
            "D. 80 cm²",
            "</CH>",
            "</Gr>",
            "",
            "<Gr> Kiến thức Tiếng Anh",
            "<CH>",
            "Câu 3. Từ nào sau đây có nghĩa là 'sách'?",
            "A. Pen",
            "*B. Book",
            "C. Table",
            "D. Chair",
            "</CH>",
            "<CH>",
            "Câu 4. Chọn câu đúng:",
            "A. He are a student",
            "*B. He is a student",
            "C. He am a student",
            "D. He be a student",
            "</CH>",
            "</Gr>",
            "",
            "<Gr> Kiến thức Khoa học",
            "<CH>",
            "Câu 5. Trái đất quay quanh mặt trời mất khoảng thời gian:",
            "A. 1 tháng",
            "B. 6 tháng",
            "*C. 1 năm",
            "D. 2 năm",
            "</CH>",
            "</Gr>"
        ]
        
        self.parse_questions()
        self.update_question_table()
        self.update_content_viewer()
        self.update_status("Đã tải dữ liệu mẫu - 5 câu hỏi từ 3 nhóm")
    
    def parse_questions(self):
        """Phân tích câu hỏi từ raw content với error handling và nhận dạng tags"""
        self.questions = []
        self.parsing_errors = []  # Lưu trữ các lỗi phân tích
        self.question_groups = []  # Lưu thông tin nhóm câu hỏi
        
        current_question = None
        current_group = None
        question_num = 0
        in_question_block = False  # Đang trong khối <CH>...</CH>
        
        for line_idx, line in enumerate(self.raw_content):
            line = line.strip()
            
            # Bỏ qua dòng trống
            if not line:
                continue
            
            # Nhận dạng tag nhóm câu hỏi <Gr> hoặc <gr>
            if line.startswith('<Gr>') or line.startswith('<gr>'):
                group_name = line[4:].strip()  # Lấy tên nhóm sau <Gr> hoặc <gr>
                current_group = {
                    'name': group_name,
                    'start_line': line_idx,
                    'questions': []
                }
                self.question_groups.append(current_group)
                print(f"📂 Phát hiện nhóm: {group_name}")
                continue
            
            # Nhận dạng tag kết thúc nhóm </Gr> hoặc </gr>
            if line.strip() == '</Gr>' or line.strip() == '</gr>':
                if current_group:
                    print(f"📂 Kết thúc nhóm: {current_group['name']} ({len(current_group['questions'])} câu hỏi)")
                current_group = None  # Đặt lại group hiện tại
                continue
            
            # Nhận dạng tag bắt đầu câu hỏi <CH>
            if line == '<CH>':
                in_question_block = True
                # Validate câu hỏi trước đó nếu có
                if current_question:
                    validation_result = self.validate_question(current_question)
                    if validation_result['valid']:
                        self.questions.append(current_question)
                        if current_group:
                            current_group['questions'].append(current_question['number'])
                    else:
                        self.parsing_errors.append({
                            'question': current_question,
                            'line_idx': line_idx,
                            'error': validation_result['error'],
                            'type': 'validation_error'
                        })
                    current_question = None
                continue
            
            # Nhận dạng tag kết thúc câu hỏi </CH>
            if line == '</CH>':
                in_question_block = False
                # Validate câu hỏi hiện tại
                if current_question:
                    validation_result = self.validate_question(current_question)
                    if validation_result['valid']:
                        self.questions.append(current_question)
                        if current_group:
                            current_group['questions'].append(current_question['number'])
                    else:
                        self.parsing_errors.append({
                            'question': current_question,
                            'line_idx': line_idx,
                            'error': validation_result['error'],
                            'type': 'validation_error'
                        })
                    current_question = None
                continue
                
            # Kiểm tra câu hỏi mới (có hoặc không có <CH>)
            if re.match(r'^Câu\s+\d+[:\.]', line):
                # Validate câu hỏi trước đó nếu có (cho format không có tag)
                if current_question and not in_question_block:
                    validation_result = self.validate_question(current_question)
                    if validation_result['valid']:
                        self.questions.append(current_question)
                        if current_group:
                            current_group['questions'].append(current_question['number'])
                    else:
                        self.parsing_errors.append({
                            'question': current_question,
                            'line_idx': line_idx,
                            'error': validation_result['error'],
                            'type': 'validation_error'
                        })
                
                question_num += 1
                # Loại bỏ phần "Câu X:" hoặc "Câu X." ở đầu
                question_text = re.sub(r'^Câu\s+\d+[:\.]', '', line).strip()
                
                current_question = {
                    'number': question_num,
                    'question': question_text,
                    'answers': {},
                    'correct_answer': None,
                    'start_line': line_idx,
                    'has_error': False,
                    'error_message': None,
                    'group': current_group['name'] if current_group else None
                }
            
            # Kiểm tra đáp án
            elif re.match(r'^[*]?[A-E]\.', line):
                if current_question:
                    try:
                        is_correct = line.startswith('*')
                        if is_correct:
                            line = line[1:]
                            if current_question['correct_answer']:
                                current_question['has_error'] = True
                                current_question['error_message'] = "Có nhiều hơn 1 đáp án đúng"
                            else:
                                current_question['correct_answer'] = line[0]
                        
                        answer_letter = line[0]
                        answer_text = line[2:].strip()
                        
                        if answer_letter in current_question['answers']:
                            current_question['has_error'] = True
                            current_question['error_message'] = f"Đáp án {answer_letter} bị trùng"
                        
                        current_question['answers'][answer_letter] = answer_text
                        
                    except Exception as e:
                        current_question['has_error'] = True
                        current_question['error_message'] = f"Lỗi phân tích đáp án: {str(e)}"
        
        # Validate câu hỏi cuối cùng
        if current_question:
            validation_result = self.validate_question(current_question)
            if validation_result['valid']:
                self.questions.append(current_question)
                if current_group:
                    current_group['questions'].append(current_question['number'])
            else:
                self.parsing_errors.append({
                    'question': current_question,
                    'line_idx': len(self.raw_content),
                    'error': validation_result['error'],
                    'type': 'validation_error'
                })
        
        print(f"Đã phân tích được {len(self.questions)} câu hỏi")
        print(f"Phát hiện {len(self.question_groups)} nhóm câu hỏi")
        
        # LƯU THỨ TỰ GỐC trước khi sắp xếp (để kiểm tra lỗi sau)
        self.original_question_order = [q['number'] for q in self.questions]
        
        # Sắp xếp câu hỏi theo thứ tự số
        self.sort_questions_by_number()
        
        if self.parsing_errors:
            print(f"⚠️  Phát hiện {len(self.parsing_errors)} lỗi cần xử lý")
            self.show_parsing_errors()
        
        print(f"📊 Tổng kết: {len(self.questions)} câu hỏi hợp lệ, {len(getattr(self, 'parsing_errors', []))} lỗi")
        print(f"📋 Thứ tự câu hỏi: {self.get_question_summary()}")
        
        # Hiển thị thông tin nhóm
        for group in self.question_groups:
            print(f"📂 Nhóm '{group['name']}': {len(group['questions'])} câu hỏi")
        
        # Cập nhật thông tin cho tab trộn đề
        if hasattr(self, 'mixer_info_label'):
            self.update_mixer_info()
    
    def update_question_table(self):
        """Cập nhật bảng câu hỏi theo format 3 cột: Label | Content | Group"""
        # Xóa dữ liệu cũ
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Nhóm câu hỏi theo group nếu có
        current_group = None
        
        # Thêm dữ liệu mới theo format: Label | Content | Group
        for i, q in enumerate(self.questions):
            question_id = f"0.{q['number']}"
            group_name = q.get('group', '')
            
            # Thêm header nhóm nếu câu hỏi thuộc nhóm mới
            if group_name and group_name != current_group:
                current_group = group_name
                group_item = self.tree.insert('', 'end', values=(
                    "",
                    f"📂 {current_group}",
                    ""
                ))
                self.tree.item(group_item, tags=('group_header',))
            
            # Thêm hàng câu hỏi với tên nhóm
            question_item = self.tree.insert('', 'end', values=(
                f"Câu {question_id}",
                q['question'],
                group_name if group_name else "-"
            ))
            self.tree.item(question_item, tags=(f"question_{i}",))
            
            # Thêm các đáp án (không hiển thị group cho đáp án)
            for letter in sorted(q['answers'].keys()):
                answer_item = self.tree.insert('', 'end', values=(
                    f"   {letter})",
                    q['answers'][letter],
                    ""
                ))
                self.tree.item(answer_item, tags=(f"answer_{i}_{letter}",))
            
            # Thêm đáp án đúng (không hiển thị group)
            correct_item = self.tree.insert('', 'end', values=(
                "   Đáp án",
                q['correct_answer'] or 'Chưa xác định',
                ""
            ))
            self.tree.item(correct_item, tags=(f"correct_{i}",))
            
            # Không thêm hàng trống để phân cách - bỏ theo yêu cầu
        
        # Configure colors cho các loại hàng khác nhau
        # Nhóm câu hỏi - màu xanh lá đậm
        self.tree.tag_configure('group_header', background='#4CAF50', foreground='white', font=('Arial', 11, 'bold'))
        
        self.tree.tag_configure('question_0', background='#e8f5e8')
        self.tree.tag_configure('question_1', background='#e8f5e8')
        self.tree.tag_configure('question_2', background='#e8f5e8')
        self.tree.tag_configure('question_3', background='#e8f5e8')
        
        for i in range(len(self.questions)):
            self.tree.tag_configure(f'question_{i}', background='#e8f5e8', font=('Arial', 10, 'bold'))
            self.tree.tag_configure(f'correct_{i}', background='#fff3cd', font=('Arial', 10, 'bold'))
            
            # Configure cho các đáp án
            for letter in ['A', 'B', 'C', 'D', 'E']:
                self.tree.tag_configure(f'answer_{i}_{letter}', background='#f8f9fa')
    
    def validate_question(self, question):
        """Validate một câu hỏi"""
        errors = []
        
        # Kiểm tra có đáp án hay không
        if not question['answers']:
            errors.append("Không có đáp án nào")
        
        # Kiểm tra có đáp án đúng hay không
        if not question['correct_answer']:
            errors.append("Không có đáp án đúng (thiếu dấu *)")
        
        # Kiểm tra đáp án đúng có tồn tại trong danh sách không
        if question['correct_answer'] and question['correct_answer'] not in question['answers']:
            errors.append(f"Đáp án đúng '{question['correct_answer']}' không tồn tại trong danh sách")
        
        # Kiểm tra số lượng đáp án (tối thiểu 2)
        if len(question['answers']) < 2:
            errors.append("Quá ít đáp án (cần ít nhất 2 đáp án)")
        
        # Kiểm tra có lỗi nội bộ không
        if question.get('has_error', False):
            errors.append(question.get('error_message', 'Lỗi không xác định'))
        
        return {
            'valid': len(errors) == 0,
            'error': '; '.join(errors) if errors else None
        }
    
    def check_data_quality(self):
        """Kiểm tra chất lượng dữ liệu: thứ tự câu hỏi, câu hỏi thiếu đáp án, câu không có tag <CH>"""
        if not self.raw_content:
            messagebox.showwarning("Cảnh báo", "Chưa có dữ liệu để kiểm tra!")
            return
        
        issues = []
        
        # 1. Kiểm tra thứ tự câu hỏi
        print("\n🔍 Kiểm tra thứ tự câu hỏi...")
        sequence_issues = self.check_question_sequence()
        issues.extend(sequence_issues)
        
        # 2. Kiểm tra câu hỏi không có tag <CH>...</CH>
        print("🔍 Kiểm tra câu hỏi thiếu tag <CH>...")
        missing_tag_issues = self.check_missing_ch_tags()
        issues.extend(missing_tag_issues)
        
        # 3. Kiểm tra câu hỏi thiếu đáp án
        print("🔍 Kiểm tra câu hỏi thiếu đáp án...")
        answer_issues = self.check_incomplete_questions()
        issues.extend(answer_issues)
        
        # 4. Kiểm tra đáp án trùng chữ cái
        print("🔍 Kiểm tra đáp án trùng chữ cái...")
        duplicate_issues = self.check_duplicate_answer_letters()
        issues.extend(duplicate_issues)
        
        # 5. Kiểm tra câu hỏi trùng lặp số
        print("🔍 Kiểm tra câu hỏi trùng lặp số...")
        duplicate_question_issues = self.check_duplicate_question_numbers()
        issues.extend(duplicate_question_issues)
        
        # Hiển thị kết quả
        self.show_quality_report(issues)
    
    def check_question_sequence(self):
        """Kiểm tra thứ tự câu hỏi có đúng không - dựa trên thứ tự GỐC trong file"""
        issues = []
        
        if not self.questions:
            return issues
        
        # Sử dụng thứ tự gốc từ file (trước khi sắp xếp)
        original_order = getattr(self, 'original_question_order', [q['number'] for q in self.questions])
        
        expected_num = 1
        for i, actual_num in enumerate(original_order):
            if actual_num != expected_num:
                # Tìm câu hỏi này để lấy thông tin dòng
                question = next((q for q in self.questions if q['number'] == actual_num), None)
                
                issues.append({
                    'type': 'Sai thứ tự',
                    'severity': 'error',
                    'message': f"Câu {actual_num} xuất hiện ở vị trí {i+1}, mong đợi câu {expected_num}",
                    'question_num': actual_num,
                    'expected': expected_num,
                    'position': i + 1,
                    'line': question.get('start_line', 0) + 1 if question else i + 1,
                    'raw_line_start': question.get('start_line', 0) if question else i
                })
                print(f"⚠️  Sai thứ tự: Vị trí {i+1} có câu {actual_num}, mong đợi câu {expected_num}")
            
            expected_num = actual_num + 1
        
        return issues
    
    def check_missing_ch_tags(self):
        """Kiểm tra câu hỏi không nằm trong tag <CH>...</CH>"""
        issues = []
        
        in_ch_block = False
        question_in_block = {}
        current_question_num = None
        
        for line_idx, line in enumerate(self.raw_content):
            line = line.strip()
            
            if line == '<CH>':
                in_ch_block = True
                current_question_num = None
                continue
            
            if line == '</CH>':
                in_ch_block = False
                if current_question_num:
                    question_in_block[current_question_num] = True
                current_question_num = None
                continue
            
            # Phát hiện câu hỏi
            match = re.match(r'^Câu\s+(\d+)[:\.]', line)
            if match:
                q_num = int(match.group(1))
                if in_ch_block:
                    current_question_num = q_num
                else:
                    # Câu hỏi không nằm trong <CH>...</CH>
                    issues.append({
                        'type': 'Thiếu tag <CH>',
                        'severity': 'warning',
                        'message': f"Câu {q_num} không nằm trong tag <CH>...</CH> (dòng {line_idx + 1})",
                        'question_num': q_num,
                        'line': line_idx + 1
                    })
                    print(f"⚠️  Thiếu tag: Câu {q_num} không có tag <CH> ở dòng {line_idx + 1}")
        
        return issues
    
    def check_incomplete_questions(self):
        """Kiểm tra câu hỏi thiếu đáp án, không đủ đáp án, và đáp án ghi chung dòng"""
        issues = []
        
        for q in self.questions:
            q_num = q['number']
            start_line = q.get('start_line', 0)
            
            # Kiểm tra số lượng đáp án
            num_answers = len(q['answers'])
            
            if num_answers == 0:
                issues.append({
                    'type': 'Không có đáp án',
                    'severity': 'error',
                    'message': f"Câu {q_num} không có đáp án nào",
                    'question_num': q_num,
                    'line': start_line + 1,
                    'raw_line_start': start_line
                })
            elif num_answers == 1:
                issues.append({
                    'type': 'Thiếu đáp án',
                    'severity': 'warning',
                    'message': f"Câu {q_num} chỉ có {num_answers} đáp án (nên có ít nhất 2)",
                    'question_num': q_num,
                    'line': start_line + 1,
                    'raw_line_start': start_line
                })
            
            # Kiểm tra đáp án đúng
            if not q['correct_answer']:
                issues.append({
                    'type': 'Thiếu đáp án đúng (*)',
                    'severity': 'error',
                    'message': f"Câu {q_num} không có đáp án nào được đánh dấu * là đáp án đúng",
                    'question_num': q_num,
                    'line': start_line + 1,
                    'raw_line_start': start_line
                })
                print(f"🔴 Câu {q_num}: Không có đáp án đúng (thiếu dấu *)")
        
        # Kiểm tra đáp án ghi chung một dòng trong raw_content
        combined_answer_issues = self.check_combined_answers()
        issues.extend(combined_answer_issues)
        
        return issues
    
    def check_combined_answers(self):
        """Kiểm tra đáp án bị ghi chung một dòng (ví dụ: 'A. Đáp án A B. Đáp án B' hoặc 'C. 24 tháng.					D. 18 tháng.')"""
        issues = []
        
        current_question_num = None
        in_question_block = False
        
        for line_idx, line in enumerate(self.raw_content):
            line_stripped = line.strip()
            
            # Phát hiện tag <CH>
            if line_stripped == '<CH>':
                in_question_block = True
                continue
            
            if line_stripped == '</CH>':
                in_question_block = False
                current_question_num = None
                continue
            
            # Phát hiện câu hỏi
            match = re.match(r'^Câu\s+(\d+)[:\.]', line_stripped)
            if match:
                current_question_num = int(match.group(1))
                continue
            
            # Kiểm tra dòng có nhiều đáp án (A. ... B. ... hoặc A. ... C. ...)
            if current_question_num and line_stripped:
                # Đếm số lượng pattern đáp án trong một dòng
                answer_patterns = re.findall(r'[*]?[A-E]\.', line_stripped)
                
                # Chỉ báo lỗi nếu có nhiều hơn 1 pattern VÀ không phải là nội dung đáp án
                # Bỏ qua nếu dòng bắt đầu bằng một đáp án duy nhất (đó là nội dung của đáp án)
                if len(answer_patterns) > 1:
                    # Kiểm tra xem có phải là nội dung của một đáp án không
                    # Nếu dòng bắt đầu bằng *X. hoặc X. thì đó là đáp án chính
                    first_answer_match = re.match(r'^([*]?[A-E]\.)', line_stripped)
                    if first_answer_match:
                        # Lấy vị trí kết thúc của pattern đầu tiên
                        first_pattern_end = first_answer_match.end()
                        # Lấy phần còn lại của dòng sau pattern đầu tiên
                        remaining_text = line_stripped[first_pattern_end:].strip()
                        
                        # CÁCH 1: Kiểm tra nếu có pattern ngay đầu phần còn lại (không phải trong nội dung)
                        has_pattern_at_start = remaining_text and re.match(r'^[*]?[A-E]\.', remaining_text)
                        
                        # CÁCH 2: Kiểm tra nếu có nhiều khoảng trắng/tab giữa các đáp án
                        # Ví dụ: "C. 24 tháng.					D. 18 tháng."
                        has_whitespace_separator = False
                        if '\t' in line or '  ' in line:  # Tab hoặc 2+ spaces liên tiếp
                            # Tìm vị trí của pattern thứ 2 trong dòng gốc (chưa strip)
                            second_pattern = re.search(r'\s{2,}[*]?[A-E]\.|\t+[*]?[A-E]\.', line)
                            if second_pattern:
                                has_whitespace_separator = True
                        
                        if has_pattern_at_start or has_whitespace_separator:
                            issues.append({
                                'type': 'Đáp án ghi chung dòng',
                                'severity': 'error',
                                'message': f"Câu {current_question_num}: Có {len(answer_patterns)} đáp án ghi chung một dòng (dòng {line_idx + 1}): '{line_stripped[:80]}...'",
                                'question_num': current_question_num,
                                'line': line_idx + 1,
                                'raw_line_start': line_idx,
                                'detail': f"Phát hiện: {', '.join(answer_patterns)}"
                            })
                            print(f"🔴 Câu {current_question_num} dòng {line_idx + 1}: {len(answer_patterns)} đáp án chung dòng - {answer_patterns}")
        
        return issues
    
    def check_duplicate_answer_letters(self):
        """Kiểm tra đáp án trùng chữ cái (ví dụ: A, *B, B, C)"""
        issues = []
        
        for q in self.questions:
            q_num = q['number']
            start_line = q.get('start_line', 0)
            
            # Lấy danh sách chữ cái của tất cả đáp án
            # answers là dict với key là chữ cái (A, B, C...)
            answer_letters = list(q['answers'].keys())
            
            # KHÔNG thêm correct_answer vì nó đã có trong answers rồi
            # (correct_answer chỉ đánh dấu đáp án nào là đúng)
            
            # Kiểm tra trùng lặp
            seen = set()
            duplicates = set()
            for letter in answer_letters:
                if letter in seen:
                    duplicates.add(letter)
                seen.add(letter)
            
            if duplicates:
                issues.append({
                    'type': 'Đáp án trùng chữ cái',
                    'severity': 'error',
                    'message': f"Câu {q_num}: Phát hiện đáp án trùng chữ cái: {', '.join(sorted(duplicates))}",
                    'question_num': q_num,
                    'line': start_line + 1,
                    'raw_line_start': start_line,
                    'detail': f"Tất cả chữ cái: {', '.join(answer_letters)}"
                })
                print(f"🔴 Câu {q_num}: Đáp án trùng lặp - {', '.join(sorted(duplicates))}")
        
        return issues
    
    def check_duplicate_question_numbers(self):
        """Kiểm tra câu hỏi có số trùng lặp (ví dụ: Câu 156 xuất hiện 2 lần)"""
        issues = []
        
        # Đếm số lần xuất hiện của mỗi số câu hỏi
        question_counts = {}
        question_positions = {}  # Lưu vị trí các câu hỏi trùng
        
        for i, q in enumerate(self.questions):
            q_num = q['number']
            
            if q_num not in question_counts:
                question_counts[q_num] = 0
                question_positions[q_num] = []
            
            question_counts[q_num] += 1
            question_positions[q_num].append({
                'position': i + 1,
                'line': q.get('start_line', 0) + 1,
                'raw_line_start': q.get('start_line', 0)
            })
        
        # Tìm các câu hỏi trùng lặp
        for q_num, count in question_counts.items():
            if count > 1:
                positions = question_positions[q_num]
                position_str = ', '.join([f"vị trí {p['position']} (dòng {p['line']})" for p in positions])
                
                # Tạo issue cho mỗi lần xuất hiện (trừ lần đầu tiên)
                for i, pos in enumerate(positions):
                    if i > 0:  # Bỏ qua lần đầu tiên
                        issues.append({
                            'type': 'Câu hỏi trùng lặp',
                            'severity': 'error',
                            'message': f"Câu {q_num} xuất hiện {count} lần: {position_str}",
                            'question_num': q_num,
                            'line': pos['line'],
                            'raw_line_start': pos['raw_line_start'],
                            'detail': f"Lần xuất hiện thứ {i + 1}/{count}"
                        })
                
                print(f"🔴 Câu {q_num} trùng lặp {count} lần tại: {position_str}")
        
        return issues
    
    def show_quality_report(self, issues):
        """Hiển thị báo cáo chất lượng dữ liệu"""
        report_window = tk.Toplevel(self.root)
        report_window.title("Báo cáo kiểm tra chất lượng dữ liệu")
        report_window.geometry("900x700")
        
        # Header
        header_frame = tk.Frame(report_window, bg='#f0f0f0', height=60)
        header_frame.pack(fill=tk.X, padx=10, pady=5)
        header_frame.pack_propagate(False)
        
        if not issues:
            tk.Label(header_frame, text="✅ Dữ liệu hoàn hảo - Không phát hiện vấn đề!", 
                    font=('Arial', 14, 'bold'), bg='#f0f0f0', fg='green').pack(pady=15)
        else:
            error_count = len([i for i in issues if i['severity'] == 'error'])
            warning_count = len([i for i in issues if i['severity'] == 'warning'])
            info_count = len([i for i in issues if i['severity'] == 'info'])
            
            tk.Label(header_frame, 
                    text=f"🔍 Phát hiện {len(issues)} vấn đề: {error_count} lỗi, {warning_count} cảnh báo, {info_count} thông tin", 
                    font=('Arial', 12, 'bold'), bg='#f0f0f0', fg='#d32f2f').pack(pady=15)
        
        # Treeview để hiển thị issues
        list_frame = tk.Frame(report_window)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        columns = ('Loại', 'Mức độ', 'Mô tả', 'Câu hỏi', 'Dòng')
        tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=20)
        
        tree.heading('Loại', text='Loại vấn đề')
        tree.heading('Mức độ', text='Mức độ')
        tree.heading('Mô tả', text='Chi tiết')
        tree.heading('Câu hỏi', text='Câu số')
        tree.heading('Dòng', text='Dòng')
        
        tree.column('Loại', width=150)
        tree.column('Mức độ', width=80)
        tree.column('Mô tả', width=450)
        tree.column('Câu hỏi', width=70)
        tree.column('Dòng', width=70)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Thêm dữ liệu và lưu reference
        issue_map = {}
        for idx, issue in enumerate(sorted(issues, key=lambda x: (x.get('question_num', 0), x['severity']))):
            severity_icon = {
                'error': '🔴',
                'warning': '⚠️',
                'info': 'ℹ️'
            }.get(issue['severity'], '')
            
            item = tree.insert('', 'end', values=(
                issue['type'],
                f"{severity_icon} {issue['severity'].upper()}",
                issue['message'],
                issue.get('question_num', '-'),
                issue.get('line', '-')
            ))
            issue_map[item] = issue
        
        # Double click để xem nguồn
        def on_double_click(event):
            selection = tree.selection()
            if selection:
                item = selection[0]
                issue = issue_map.get(item)
                if issue and 'raw_line_start' in issue:
                    self.show_raw_content_at_line(issue['raw_line_start'], issue.get('question_num'))
        
        tree.bind('<Double-1>', on_double_click)
        
        # Thêm label hướng dẫn
        instruction_frame = tk.Frame(report_window, bg='#e3f2fd', height=30)
        instruction_frame.pack(fill=tk.X, padx=10, pady=(0, 5))
        instruction_frame.pack_propagate(False)
        
        tk.Label(instruction_frame, 
                text="💡 Nhấp đúp vào lỗi để xem vị trí trong file gốc", 
                bg='#e3f2fd', fg='#1976d2', font=('Arial', 9, 'italic')).pack(pady=5)
        
        # Buttons
        button_frame = tk.Frame(report_window)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def export_report():
            """Xuất báo cáo ra file text"""
            file_path = filedialog.asksaveasfilename(
                title="Lưu báo cáo",
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("=" * 80 + "\n")
                    f.write("BÁO CÁO KIỂM TRA CHẤT LƯỢNG DỮ LIỆU\n")
                    f.write("=" * 80 + "\n\n")
                    
                    if not issues:
                        f.write("✅ Dữ liệu hoàn hảo - Không phát hiện vấn đề!\n")
                    else:
                        error_count = len([i for i in issues if i['severity'] == 'error'])
                        warning_count = len([i for i in issues if i['severity'] == 'warning'])
                        info_count = len([i for i in issues if i['severity'] == 'info'])
                        
                        f.write(f"Tổng số vấn đề: {len(issues)}\n")
                        f.write(f"  - Lỗi: {error_count}\n")
                        f.write(f"  - Cảnh báo: {warning_count}\n")
                        f.write(f"  - Thông tin: {info_count}\n\n")
                        f.write("=" * 80 + "\n\n")
                        
                        for issue in sorted(issues, key=lambda x: (x.get('question_num', 0), x['severity'])):
                            f.write(f"[{issue['severity'].upper()}] {issue['type']}\n")
                            f.write(f"  Câu hỏi: {issue.get('question_num', 'N/A')}\n")
                            f.write(f"  Chi tiết: {issue['message']}\n")
                            f.write("-" * 80 + "\n")
                
                messagebox.showinfo("Thành công", f"Đã xuất báo cáo: {os.path.basename(file_path)}")
        
        tk.Button(button_frame, text="📄 Xuất báo cáo", command=export_report,
                 bg='#2196F3', fg='white', padx=20).pack(side=tk.LEFT, padx=5)
        
        tk.Button(button_frame, text="Đóng", command=report_window.destroy,
                 bg='#6c757d', fg='white', padx=20).pack(side=tk.RIGHT, padx=5)
    
    def show_raw_content_at_line(self, line_num, question_num=None):
        """Hiển thị raw content tại dòng chỉ định và highlight
        
        Args:
            line_num: Index của dòng trong raw_content (0-based)
            question_num: Số câu hỏi (optional)
        """
        # Focus vào panel raw content
        self.content_text.focus_set()
        
        # Chuyển đổi từ index (0-based) sang line number trong Text widget (1-based)
        actual_line = line_num + 1
        
        # Xóa highlight cũ
        self.content_text.tag_remove('error_highlight', '1.0', tk.END)
        
        # Highlight dòng lỗi với màu đỏ nhạt
        # Text widget: dòng 1 = "1.0", dòng 2 = "2.0", ...
        start_pos = f"{actual_line}.0"
        end_pos = f"{actual_line}.end"
        self.content_text.tag_add('error_highlight', start_pos, end_pos)
        self.content_text.tag_configure('error_highlight', background='#ffcccc', font=('Consolas', 11, 'bold'))
        
        # Đẩy dòng lỗi lên gần đầu màn hình (dòng thứ 3 từ trên xuống)
        # Bước 1: Scroll đến dòng trước đó vài dòng
        context_lines_before = 2  # Hiển thị 2 dòng context phía trên
        scroll_to_line = max(1, actual_line - context_lines_before)
        
        # Bước 2: Scroll đến vị trí đó và đặt nó ở đầu view
        self.content_text.see(f"{scroll_to_line}.0")
        
        # Bước 3: Đảm bảo dòng lỗi được nhìn thấy
        self.content_text.see(start_pos)
        
        # Cập nhật status
        msg = f"🔍 Đang xem dòng {actual_line}"
        if question_num:
            msg += f" (Câu {question_num})"
        self.update_status(msg)
        
        print(f"📍 Focus vào dòng {actual_line}" + (f" - Câu {question_num}" if question_num else ""))
    
    def show_parsing_errors(self):
        """Hiển thị dialog với các lỗi phân tích"""
        error_window = tk.Toplevel(self.root)
        error_window.title(f"Lỗi phân tích câu hỏi - Tổng {len(self.parsing_errors)} lỗi")
        error_window.geometry("800x600")
        error_window.grab_set()  # Modal dialog
        
        # Header
        header_frame = tk.Frame(error_window, bg='#ffcccc', height=50)
        header_frame.pack(fill=tk.X, padx=10, pady=5)
        header_frame.pack_propagate(False)
        
        tk.Label(header_frame, text=f"⚠️ Phát hiện {len(self.parsing_errors)} lỗi khi phân tích câu hỏi - Vui lòng sửa từng lỗi", 
                font=('Arial', 12, 'bold'), bg='#ffcccc', fg='darkred').pack(pady=10)
        
        # Error list
        list_frame = tk.Frame(error_window)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Treeview để hiển thị lỗi
        columns = ('Câu', 'Lỗi', 'Hành động')
        error_tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=10)
        
        error_tree.heading('Câu', text='Câu hỏi')
        error_tree.heading('Lỗi', text='Mô tả lỗi')
        error_tree.heading('Hành động', text='Hành động')
        
        error_tree.column('Câu', width=100)
        error_tree.column('Lỗi', width=400)
        error_tree.column('Hành động', width=200)
        
        # Scrollbar
        error_scroll = ttk.Scrollbar(list_frame, orient="vertical", command=error_tree.yview)
        error_tree.configure(yscrollcommand=error_scroll.set)
        
        error_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        error_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Thêm dữ liệu lỗi
        for i, error_info in enumerate(self.parsing_errors):
            question = error_info['question']
            item = error_tree.insert('', 'end', values=(
                f"Câu {question['number']}",
                error_info['error'],
                "Nhấp đúp để sửa"
            ))
            # Lưu reference
            error_tree.item(item, tags=(str(i),))
        
        # Bind double click
        def on_error_double_click(event):
            selection = error_tree.selection()
            if selection:
                item = selection[0]
                tags = error_tree.item(item, 'tags')
                if tags:
                    error_idx = int(tags[0])
                    self.manual_edit_question(error_idx, error_window)
        
        error_tree.bind('<Double-1>', on_error_double_click)
        
        # Buttons
        button_frame = tk.Frame(error_window)
        button_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Button(button_frame, text="Đóng", command=error_window.destroy,
                 bg='#6c757d', fg='white', padx=20).pack(side=tk.RIGHT, padx=5)
        
        tk.Button(button_frame, text="Bỏ qua tất cả lỗi", 
                 command=lambda: self.ignore_all_errors(error_window),
                 bg='#dc3545', fg='white', padx=20).pack(side=tk.RIGHT, padx=5)
    
    def manual_edit_question(self, error_idx, parent_window=None):
        """Mở dialog để chỉnh sửa thủ công câu hỏi bị lỗi"""
        if error_idx >= len(self.parsing_errors):
            return
        
        error_info = self.parsing_errors[error_idx]
        question = error_info['question']
        
        # Tạo dialog chỉnh sửa
        edit_window = tk.Toplevel(parent_window or self.root)
        edit_window.title(f"Chỉnh sửa Câu {question['number']}")
        edit_window.geometry("700x500")
        edit_window.grab_set()
        
        # Question text
        tk.Label(edit_window, text="Nội dung câu hỏi:", font=('Arial', 10, 'bold')).pack(anchor='w', padx=10, pady=(10,5))
        question_text = scrolledtext.ScrolledText(edit_window, height=3, wrap=tk.WORD)
        question_text.pack(fill=tk.X, padx=10, pady=5)
        question_text.insert(1.0, question['question'])
        
        # Answers
        tk.Label(edit_window, text="Đáp án:", font=('Arial', 10, 'bold')).pack(anchor='w', padx=10, pady=(10,5))
        
        answer_frame = tk.Frame(edit_window)
        answer_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        answer_texts = {}
        correct_var = tk.StringVar(value=question.get('correct_answer', ''))
        
        # Tạo các ô nhập đáp án
        for letter in ['A', 'B', 'C', 'D', 'E']:
            row_frame = tk.Frame(answer_frame)
            row_frame.pack(fill=tk.X, pady=2)
            
            # Radio button cho đáp án đúng
            tk.Radiobutton(row_frame, variable=correct_var, value=letter, 
                          text=f"{letter}:", font=('Arial', 10, 'bold')).pack(side=tk.LEFT)
            
            # Text entry cho nội dung đáp án
            entry = tk.Entry(row_frame, font=('Arial', 10))
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5,0))
            
            # Điền dữ liệu có sẵn
            if letter in question['answers']:
                entry.insert(0, question['answers'][letter])
            
            answer_texts[letter] = entry
        
        # Buttons
        button_frame = tk.Frame(edit_window)
        button_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def save_question():
            # Thu thập dữ liệu
            new_question = question_text.get(1.0, tk.END).strip()
            new_answers = {}
            new_correct = correct_var.get()
            
            for letter, entry in answer_texts.items():
                content = entry.get().strip()
                if content:
                    new_answers[letter] = content
            
            # Validate
            if not new_question:
                messagebox.showerror("Lỗi", "Vui lòng nhập nội dung câu hỏi!")
                return
            
            if len(new_answers) < 2:
                messagebox.showerror("Lỗi", "Cần có ít nhất 2 đáp án!")
                return
            
            if not new_correct or new_correct not in new_answers:
                messagebox.showerror("Lỗi", "Vui lòng chọn đáp án đúng từ các đáp án đã nhập!")
                return
            
            # Cập nhật câu hỏi
            question['question'] = new_question
            question['answers'] = new_answers
            question['correct_answer'] = new_correct
            question['has_error'] = False
            question['error_message'] = None
            
            # Tìm vị trí đúng để chèn câu hỏi theo thứ tự số
            insert_position = 0
            for i, existing_q in enumerate(self.questions):
                if existing_q['number'] > question['number']:
                    insert_position = i
                    break
                insert_position = i + 1
            
            # Chèn câu hỏi vào đúng vị trí thứ tự
            self.questions.insert(insert_position, question)
            
            # Xóa khỏi danh sách lỗi
            self.parsing_errors.pop(error_idx)
            
            # Cập nhật status
            total_errors = len(self.parsing_errors)
            valid_questions = len([q for q in self.questions if not q.get('has_error', False)])
            status_msg = f"Đã sửa lỗi Câu {question['number']} - Tổng {len(self.questions)} câu hỏi hợp lệ (đã sắp xếp) | {valid_questions} câu hợp lệ, {total_errors} lỗi"
            self.update_status(status_msg)
            print(f"📌 Status: {status_msg}")
            
            messagebox.showinfo("Thành công", 
                              f"Đã lưu Câu {question['number']} vào vị trí thứ {insert_position + 1}\n\n"
                              f"Còn {total_errors} lỗi cần sửa.")
            
            # Đóng cửa sổ chỉnh sửa
            edit_window.destroy()
            
            # KHÔNG đóng cửa sổ danh sách lỗi, thay vào đó cập nhật lại danh sách
            if parent_window and total_errors > 0:
                # Refresh danh sách lỗi
                self.refresh_error_window(parent_window)
            elif parent_window and total_errors == 0:
                # Nếu không còn lỗi nào, đóng cửa sổ lỗi
                messagebox.showinfo("Hoàn thành", "Đã sửa hết tất cả các lỗi!")
                parent_window.destroy()
            
            # Cập nhật giao diện ngay lập tức
            self.update_question_table()
            self.update_content_viewer()
        
        tk.Button(button_frame, text="Lưu", command=save_question,
                 bg='#28a745', fg='white', padx=20).pack(side=tk.RIGHT, padx=5)
        
        tk.Button(button_frame, text="Hủy", command=edit_window.destroy,
                 bg='#6c757d', fg='white', padx=20).pack(side=tk.RIGHT, padx=5)
    
    def refresh_error_window(self, error_window):
        """Cập nhật lại danh sách lỗi trong cửa sổ"""
        # Tìm error_tree widget trong cửa sổ
        for widget in error_window.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    if isinstance(child, ttk.Treeview):
                        error_tree = child
                        # Xóa tất cả items cũ
                        for item in error_tree.get_children():
                            error_tree.delete(item)
                        
                        # Thêm lại danh sách lỗi mới
                        for i, error_info in enumerate(self.parsing_errors):
                            question = error_info['question']
                            item = error_tree.insert('', 'end', values=(
                                f"Câu {question['number']}",
                                error_info['error'],
                                "Nhấp đúp để sửa"
                            ))
                            # Lưu reference
                            error_tree.item(item, tags=(str(i),))
                        
                        # Cập nhật title
                        error_window.title(f"Lỗi phân tích câu hỏi - Còn {len(self.parsing_errors)} lỗi")
                        break
    
    def ignore_all_errors(self, error_window):
        """Bỏ qua tất cả lỗi và tiếp tục"""
        result = messagebox.askyesno("Xác nhận", 
                                   f"Bạn có chắc muốn bỏ qua {len(self.parsing_errors)} lỗi?\n"
                                   "Các câu hỏi bị lỗi sẽ không được đưa vào kết quả cuối.")
        if result:
            self.parsing_errors.clear()
            error_window.destroy()
            
            # Sắp xếp lại danh sách câu hỏi theo thứ tự
            self.sort_questions_by_number()
            
            # Cập nhật giao diện ngay lập tức thay vì refresh toàn bộ
            self.update_question_table()
            self.update_content_viewer()
            self.update_status(f"Đã bỏ qua lỗi - Tổng {len(self.questions)} câu hỏi hợp lệ (đã sắp xếp)")
            
            print(f"📋 Thứ tự sau khi sắp xếp: {self.get_question_summary()}")
    
    def sort_questions_by_number(self):
        """Sắp xếp danh sách câu hỏi theo số thứ tự"""
        self.questions.sort(key=lambda q: q['number'])
        print(f"📋 Đã sắp xếp {len(self.questions)} câu hỏi theo thứ tự số")
    
    def get_question_summary(self):
        """Lấy tóm tắt danh sách câu hỏi để debug"""
        if not self.questions:
            return "Không có câu hỏi nào"
        
        summary = []
        for i, q in enumerate(self.questions):
            summary.append(f"Vị trí {i+1}: Câu {q['number']}")
        
        return " | ".join(summary)
    
    def update_content_viewer(self):
        """Cập nhật viewer nội dung file gốc"""
        self.content_text.delete(1.0, tk.END)
        
        content = '\n'.join(self.raw_content)
        self.content_text.insert(1.0, content)
        
        # Cập nhật line numbers
        self.update_line_numbers()
        
        # Apply syntax highlighting
        self.apply_syntax_highlighting()
    
    def update_line_numbers(self):
        """Cập nhật số dòng"""
        # Đếm số dòng trong content
        line_count = int(self.content_text.index('end-1c').split('.')[0])
        
        # Tạo chuỗi số dòng
        line_numbers_text = '\n'.join(str(i) for i in range(1, line_count + 1))
        
        # Cập nhật line numbers widget
        self.line_numbers.config(state='normal')
        self.line_numbers.delete(1.0, tk.END)
        self.line_numbers.insert(1.0, line_numbers_text)
        self.line_numbers.config(state='disabled')
    
    def apply_syntax_highlighting(self):
        """Áp dụng highlight cho nội dung với error detection"""
        content = self.content_text.get(1.0, tk.END)
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            line_num = i + 1
            start_pos = f"{line_num}.0"
            end_pos = f"{line_num}.end"
            
            # Highlight câu hỏi
            if re.match(r'^Câu\s+\d+\.', line.strip()):
                self.content_text.tag_add('question', start_pos, end_pos)
                
                # Kiểm tra xem câu hỏi này có bị lỗi không
                question_num = re.search(r'Câu\s+(\d+)', line)
                if question_num:
                    num = int(question_num.group(1))
                    # Tìm trong danh sách lỗi
                    for error_info in getattr(self, 'parsing_errors', []):
                        if error_info['question']['number'] == num:
                            # Highlight màu đỏ cho câu hỏi lỗi
                            self.content_text.tag_add('error', start_pos, end_pos)
                            break
            
            # Highlight đáp án đúng
            elif line.strip().startswith('*'):
                self.content_text.tag_add('correct_answer', start_pos, end_pos)
        
        # Thêm legend/tooltip cho user hiểu các màu
        self.update_highlighting_legend()
    
    def update_highlighting_legend(self):
        """Cập nhật legend giải thích các màu highlight"""
        if hasattr(self, 'parsing_errors') and self.parsing_errors:
            legend_text = "🔴 Đỏ: Câu hỏi có lỗi | 🔵 Xanh: Câu hỏi bình thường | 🟢 Xanh lá: Đáp án đúng | 🟡 Vàng: Đang chọn"
        else:
            legend_text = "🔵 Xanh: Câu hỏi | 🟢 Xanh lá: Đáp án đúng | 🟡 Vàng: Đang chọn"
        
        # Cập nhật trong status hoặc tạo một label riêng
        if hasattr(self, 'legend_label'):
            self.legend_label.config(text=legend_text)
        else:
            # Tạo legend label nếu chưa có
            legend_frame = tk.Frame(self.root, bg='#f8f9fa', height=25)
            legend_frame.pack(fill=tk.X, side=tk.BOTTOM, before=self.status_label.master)
            legend_frame.pack_propagate(False)
            
            self.legend_label = tk.Label(legend_frame, text=legend_text, 
                                       bg='#f8f9fa', fg='#666', font=('Arial', 8))
            self.legend_label.pack(pady=3)
    
    def on_question_select(self, event):
        """Xử lý khi chọn hàng trong bảng"""
        selection = self.tree.selection()
        if not selection:
            return
        
        item = selection[0]
        # Lấy tags để xác định loại hàng được chọn
        tags = self.tree.item(item, 'tags')
        if not tags:
            return
        
        tag = tags[0]
        
        # Xử lý các loại hàng khác nhau
        if tag.startswith('question_'):
            # Được chọn là câu hỏi
            question_idx = int(tag.split('_')[1])
            if question_idx < len(self.questions):
                question = self.questions[question_idx]
                self.highlight_question_in_content(question)
                self.update_status(f"Đang xem Câu 0.{question['number']}: {question['question'][:50]}...")
        
        elif tag.startswith('answer_'):
            # Được chọn là đáp án
            parts = tag.split('_')
            question_idx = int(parts[1])
            answer_letter = parts[2]
            if question_idx < len(self.questions):
                question = self.questions[question_idx]
                self.highlight_question_in_content(question)
                answer_text = question['answers'].get(answer_letter, '')
                self.update_status(f"Đang xem Câu 0.{question['number']} - Đáp án {answer_letter}: {answer_text[:30]}...")
        
        elif tag.startswith('correct_'):
            # Được chọn là đáp án đúng
            question_idx = int(tag.split('_')[1])
            if question_idx < len(self.questions):
                question = self.questions[question_idx]
                self.highlight_question_in_content(question)
                correct_answer = question.get('correct_answer', 'Chưa xác định')
                self.update_status(f"Đang xem Câu 0.{question['number']} - Đáp án đúng: {correct_answer}")
    
    def highlight_question_in_content(self, question):
        """Highlight câu hỏi được chọn trong content viewer với hiệu ứng bôi vàng"""
        # Xóa tất cả highlight cũ
        self.content_text.tag_remove('highlight', 1.0, tk.END)
        self.content_text.tag_remove('selected_question', 1.0, tk.END)
        
        # Configure tag mới cho selection với hiệu ứng đẹp hơn
        self.content_text.tag_configure('selected_question', 
                                      background='#ffeb3b', 
                                      foreground='black', 
                                      font=('Consolas', 11, 'bold'),
                                      relief='raised',
                                      borderwidth=1)
        
        # Tìm và highlight câu hỏi
        content = self.content_text.get(1.0, tk.END)
        lines = content.split('\n')
        
        question_found = False
        highlight_start = None
        highlight_end = None
        
        # Tìm bằng pattern "Câu X:" hoặc "Câu X." trong raw content
        question_pattern = f"Câu {question['number']}"
        
        for i, line in enumerate(lines):
            line_num = i + 1
            
            # Tìm câu hỏi bắt đầu bằng pattern "Câu X:"
            if re.match(rf'^\s*Câu\s+{question["number"]}[:\.]', line.strip()):
                question_found = True
                highlight_start = line_num
                print(f"Tìm thấy câu hỏi {question['number']} tại dòng {line_num}")
            
            # Tìm kết thúc câu hỏi (câu hỏi tiếp theo hoặc end)
            elif question_found:
                # Kiểm tra nếu là câu hỏi mới khác
                if re.match(r'^Câu\s+\d+[:\.]', line.strip()) and not re.match(rf'^\s*Câu\s+{question["number"]}[:\.]', line.strip()):
                    highlight_end = line_num - 1
                    break
                # Nếu đến cuối file
                elif line_num == len(lines):
                    highlight_end = line_num
                    break
        
        # Nếu không tìm thấy câu hỏi tiếp theo, highlight đến cuối
        if highlight_start and not highlight_end:
            highlight_end = len(lines)
        
        # Apply highlight với hiệu ứng bôi vàng
        if highlight_start and highlight_end:
            start_pos = f"{highlight_start}.0"
            end_pos = f"{highlight_end}.end"
            
            # Highlight toàn bộ vùng câu hỏi
            self.content_text.tag_add('selected_question', start_pos, end_pos)
            
            # Focus và scroll đến vị trí với animation mượt
            self.content_text.focus_set()
            self.content_text.see(start_pos)
            
            # Scroll thêm để đảm bảo hiển thị đầy đủ
            try:
                self.content_text.mark_set('insert', start_pos)
                self.content_text.see('insert')
                
                # Scroll up một chút để hiển thị context
                current_line = float(start_pos.split('.')[0])
                if current_line > 3:
                    scroll_pos = f"{int(current_line - 2)}.0"
                    self.content_text.see(scroll_pos)
                    
            except Exception as e:
                print(f"Lỗi khi scroll: {e}")
            
            # Cập nhật status với thông tin chi tiết
            self.update_status(f"🔍 Đang xem Câu {question['number']} (dòng {highlight_start}-{highlight_end}): {question['question'][:40]}...")
            
            print(f"Đã highlight từ dòng {highlight_start} đến {highlight_end}")
        else:
            # Nếu không tìm thấy, thông báo lỗi
            self.update_status(f"⚠️ Không tìm thấy Câu {question['number']} trong nội dung")
            print(f"Không tìm thấy câu hỏi: {question['question'][:50]}...")
    
    def open_file(self):
        """Mở file .doc/.docx/.txt/.xml"""
        file_path = filedialog.askopenfilename(
            title="Chọn file câu hỏi",
            filetypes=[
                ("All supported", "*.doc *.docx *.txt *.xml"),
                ("Word files", "*.doc *.docx"),
                ("Text files", "*.txt"),
                ("XML files", "*.xml"),
                ("All files", "*.*")
            ]
        )
        
        if file_path:
            try:
                self.current_file = file_path
                self.load_file_content(file_path)
                self.update_status(f"Đã mở file: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể mở file: {str(e)}")
    
    def load_file_content(self, file_path):
        """Load nội dung từ file Word, TXT hoặc XML"""
        try:
            if file_path.endswith('.docx'):
                # Đọc file Word
                doc = Document(file_path)
                self.raw_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        self.raw_content.append(paragraph.text.strip())
            
            elif file_path.endswith('.txt'):
                # Đọc file TXT
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.raw_content = [line.strip() for line in f if line.strip()]
            
            elif file_path.endswith('.xml'):
                # Đọc file XML
                self.raw_content = self.parse_xml_file(file_path)
            
            else:
                messagebox.showwarning("Cảnh báo", 
                                     "File .doc cần được chuyển thành .docx trước.\n"
                                     "Hoặc sử dụng file .txt hoặc .xml")
                return
            
            self.parse_questions()
            self.update_question_table()
            self.update_content_viewer()
        
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file: {str(e)}")
    
    def parse_xml_file(self, file_path):
        """Parse file XML và trích xuất nội dung"""
        content = []
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Phương pháp 1: Nếu XML có cấu trúc chuẩn với tag <line> hoặc <content>
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    # Thêm tag đặc biệt nếu có
                    if elem.tag == 'group' or elem.tag == 'Gr':
                        content.append(f"<Gr> {elem.text.strip()}")
                    elif elem.tag == 'question' or elem.tag == 'CH':
                        content.append(f"<CH>")
                        # Lấy nội dung câu hỏi
                        if elem.text:
                            content.append(elem.text.strip())
                        # Lấy các sub-elements (đáp án, v.v.)
                        for child in elem:
                            if child.text and child.text.strip():
                                content.append(child.text.strip())
                            if child.tail and child.tail.strip():
                                content.append(child.tail.strip())
                        content.append(f"</CH>")
                    elif elem.tag not in ['root', 'questions', 'document']:
                        # Các tag khác chỉ lấy text
                        content.append(elem.text.strip())
            
            # Phương pháp 2: Nếu XML đơn giản, lấy toàn bộ text
            if not content:
                text_content = ET.tostring(root, encoding='unicode', method='text')
                content = [line.strip() for line in text_content.split('\n') if line.strip()]
            
            return content
        
        except ET.ParseError as e:
            messagebox.showerror("Lỗi XML", f"File XML không hợp lệ: {str(e)}")
            return []
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể parse XML: {str(e)}")
            return []
    
    def export_to_word(self):
        """Xuất ra file Word - Mỗi nhóm thành một file riêng"""
        if not self.questions:
            messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
            return
        
        # Kiểm tra xem có nhóm câu hỏi không
        has_groups = any(q.get('group') for q in self.questions)
        
        if has_groups:
            # Hỏi người dùng muốn xuất cách nào
            choice = messagebox.askyesnocancel(
                "Chọn cách xuất",
                "Có phát hiện các nhóm câu hỏi.\n\n"
                "- Chọn YES: Mỗi nhóm xuất thành file riêng\n"
                "- Chọn NO: Tất cả vào 1 file\n"
                "- Chọn CANCEL: Hủy"
            )
            
            if choice is None:  # Cancel
                return
            elif choice:  # Yes - Xuất từng nhóm riêng
                self.export_groups_separately()
            else:  # No - Xuất tất cả vào 1 file
                self.export_single_word_file()
        else:
            # Không có nhóm, xuất bình thường
            self.export_single_word_file()
    
    def export_single_word_file(self):
        """Xuất tất cả câu hỏi vào 1 file Word"""
        file_path = filedialog.asksaveasfilename(
            title="Lưu file kết quả",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                # Kiểm tra xem có groups không để quyết định dùng group_prefix
                has_groups = any(q.get('group') for q in self.questions)
                self.create_word_output(file_path, group_prefix=has_groups)
                messagebox.showinfo("Thành công", f"Đã xuất file: {os.path.basename(file_path)}")
                self.update_status(f"Đã xuất: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xuất file: {str(e)}")
    
    def export_groups_separately(self):
        """Xuất mỗi nhóm câu hỏi thành file Word riêng"""
        # Chọn thư mục để lưu các file
        folder_path = filedialog.askdirectory(
            title="Chọn thư mục để lưu các file Word"
        )
        
        if not folder_path:
            return
        
        try:
            # Phân nhóm câu hỏi
            groups = {}
            questions_without_group = []
            
            for q in self.questions:
                group_name = q.get('group')
                if group_name:
                    if group_name not in groups:
                        groups[group_name] = []
                    groups[group_name].append(q)
                else:
                    questions_without_group.append(q)
            
            # Xuất từng nhóm
            exported_files = []
            
            for group_name, group_questions in groups.items():
                # Tạo tên file từ tên nhóm
                safe_name = self.sanitize_filename(group_name)
                file_path = os.path.join(folder_path, f"{safe_name}.docx")
                
                # Tạo file Word cho nhóm này với renumber=True để đánh số lại từ 1
                self.create_word_output(file_path, group_questions, group_name, renumber=True)
                exported_files.append(os.path.basename(file_path))
            
            # Xuất câu hỏi không có nhóm (nếu có) với renumber=True
            if questions_without_group:
                file_path = os.path.join(folder_path, "CauHoi_KhongCoNhom.docx")
                self.create_word_output(file_path, questions_without_group, "Câu hỏi khác", renumber=True)
                exported_files.append(os.path.basename(file_path))
            
            # Thông báo thành công
            messagebox.showinfo(
                "Thành công", 
                f"Đã xuất {len(exported_files)} file:\n\n" + 
                "\n".join(f"- {name}" for name in exported_files[:10]) +
                (f"\n... và {len(exported_files) - 10} file khác" if len(exported_files) > 10 else "")
            )
            self.update_status(f"Đã xuất {len(exported_files)} file Word theo nhóm")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể xuất các file: {str(e)}")
    
    def sanitize_filename(self, name):
        """Chuyển tên nhóm thành tên file hợp lệ"""
        # Loại bỏ ký tự không hợp lệ trong tên file
        invalid_chars = '<>:"/\\|?*'
        safe_name = name
        for char in invalid_chars:
            safe_name = safe_name.replace(char, '_')
        
        # Giới hạn độ dài tên file
        if len(safe_name) > 100:
            safe_name = safe_name[:100]
        
        return safe_name.strip()
    
    def export_to_json(self):
        """Xuất ra file JSON theo format yêu cầu"""
        if not self.questions:
            messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Lưu file JSON",
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                self.create_json_output(file_path)
                messagebox.showinfo("Thành công", f"Đã xuất file JSON: {os.path.basename(file_path)}")
                self.update_status(f"Đã xuất JSON: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xuất file JSON: {str(e)}")
    
    def create_json_output(self, file_path):
        """Tạo file JSON theo format yêu cầu"""
        import json
        
        # Cấu trúc JSON theo yêu cầu
        json_data = {
            "title": "Trắc nghiệm Tin học cơ bản",
            "questions": []
        }
        
        for q in self.questions:
            question_data = {
                "id": f"0.{q['number']}",
                "question_text": q['question'],  # Nội dung câu hỏi đã được loại bỏ "Câu X:" khi parse
                "options": [],
                "correct_answer_key": q['correct_answer'] or ""
            }
            
            # Thêm các đáp án
            for letter in sorted(q['answers'].keys()):
                option = {
                    "key": letter,
                    "text": q['answers'][letter]
                }
                question_data["options"].append(option)
            
            json_data["questions"].append(question_data)
        
        # Ghi file JSON với format đẹp
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
    
    def export_to_txt(self):
        """Xuất ra file TXT"""
        if not self.questions:
            messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Lưu file TXT",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                self.create_txt_output(file_path)
                messagebox.showinfo("Thành công", f"Đã xuất file TXT: {os.path.basename(file_path)}")
                self.update_status(f"Đã xuất TXT: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xuất file TXT: {str(e)}")
    
    def export_to_xml(self):
        """Xuất ra file XML"""
        if not self.questions:
            messagebox.showwarning("Cảnh báo", "Không có dữ liệu để xuất!")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Lưu file XML",
            defaultextension=".xml",
            filetypes=[("XML files", "*.xml"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                self.create_xml_output(file_path)
                messagebox.showinfo("Thành công", f"Đã xuất file XML: {os.path.basename(file_path)}")
                self.update_status(f"Đã xuất XML: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xuất file XML: {str(e)}")
    
    def create_txt_output(self, file_path):
        """Tạo file TXT output"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("TRẮC NGHIỆM TIN HỌC CƠ BẢN\n")
            f.write("=" * 80 + "\n\n")
            
            current_group = None
            for q in self.questions:
                # Thêm header nhóm nếu có
                if q.get('group') and q.get('group') != current_group:
                    current_group = q.get('group')
                    f.write("\n" + "=" * 80 + "\n")
                    f.write(f"📂 {current_group}\n")
                    f.write("=" * 80 + "\n\n")
                
                # Câu hỏi
                question_id = f"0.{q['number']}"
                f.write(f"Câu {question_id}: {q['question']}\n")
                
                # Các đáp án
                for letter in sorted(q['answers'].keys()):
                    marker = "✓" if letter == q['correct_answer'] else " "
                    f.write(f"  [{marker}] {letter}. {q['answers'][letter]}\n")
                
                # Đáp án đúng
                f.write(f"  Đáp án: {q['correct_answer'] or 'Chưa xác định'}\n")
                f.write("\n" + "-" * 80 + "\n\n")
    
    def create_xml_output(self, file_path):
        """Tạo file XML output"""
        root = ET.Element('questions')
        root.set('title', 'Trắc nghiệm Tin học cơ bản')
        
        current_group = None
        group_elem = None
        
        for q in self.questions:
            # Tạo group element nếu có nhóm mới
            if q.get('group') and q.get('group') != current_group:
                current_group = q.get('group')
                group_elem = ET.SubElement(root, 'group')
                group_elem.set('name', current_group)
            
            # Tạo question element
            parent = group_elem if group_elem is not None else root
            question_elem = ET.SubElement(parent, 'question')
            question_elem.set('id', f"0.{q['number']}")
            
            # Nội dung câu hỏi
            question_text = ET.SubElement(question_elem, 'text')
            question_text.text = q['question']
            
            # Các đáp án
            options_elem = ET.SubElement(question_elem, 'options')
            for letter in sorted(q['answers'].keys()):
                option_elem = ET.SubElement(options_elem, 'option')
                option_elem.set('key', letter)
                option_elem.set('correct', 'true' if letter == q['correct_answer'] else 'false')
                option_elem.text = q['answers'][letter]
            
            # Đáp án đúng
            correct_elem = ET.SubElement(question_elem, 'correct_answer')
            correct_elem.text = q['correct_answer'] or ''
        
        # Tạo cây XML đẹp
        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ")  # Python 3.9+
        
        # Ghi file
        tree.write(file_path, encoding='utf-8', xml_declaration=True)
    
    def create_word_output(self, file_path, questions=None, title=None, renumber=False, group_prefix=False):
        """Tạo file Word output theo format 2 cột JSON
        
        Args:
            file_path: Đường dẫn file Word output
            questions: Danh sách câu hỏi cần xuất (None = dùng self.questions)
            title: Tiêu đề tài liệu
            renumber: True = đánh số lại từ 1, False = giữ số gốc
            group_prefix: True = đánh số theo group (Group 1: 0.x, Group 2: 1.x, ...), False = tất cả dùng 0.x
        """
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Sử dụng questions được truyền vào hoặc self.questions
        questions_to_export = questions if questions is not None else self.questions
        title_text = title if title else 'Trắc nghiệm Tin học cơ bản'
        
        doc = Document()
        doc.add_heading(title_text, 0)
        
        # Tạo bảng với 2 cột theo format JSON
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header (ẩn header vì không cần)
        header_cells = table.rows[0].cells
        header_cells[0].text = ''
        header_cells[1].text = 'Nội dung'
        
        # Format header
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.bold = True
        
        # Nếu group_prefix=True, tạo map từ group name sang group index
        group_index_map = {}
        if group_prefix:
            # Lấy danh sách các group duy nhất theo thứ tự xuất hiện
            seen_groups = []
            for q in questions_to_export:
                group_name = q.get('group', 'Không có nhóm')
                if group_name not in seen_groups:
                    seen_groups.append(group_name)
            # Tạo map: group name -> index (0, 1, 2, ...)
            for idx, group_name in enumerate(seen_groups):
                group_index_map[group_name] = idx
        
        # Đếm số câu hỏi trong mỗi group để đánh số tuần tự
        group_counters = {}
        
        # Thêm dữ liệu theo format JSON
        for idx, q in enumerate(questions_to_export, start=1):
            # Sử dụng số thứ tự mới nếu renumber=True, giữ số gốc nếu False
            question_number = idx if renumber else q['number']
            
            # Xác định prefix dựa trên group nếu group_prefix=True
            if group_prefix:
                group_name = q.get('group', 'Không có nhóm')
                group_idx = group_index_map.get(group_name, 0)
                
                # Đếm số câu trong group này
                if group_name not in group_counters:
                    group_counters[group_name] = 1
                else:
                    group_counters[group_name] += 1
                
                question_id = f"{group_idx}.{group_counters[group_name]}"
            else:
                question_id = f"0.{question_number}"
            
            # Thêm hàng câu hỏi
            row_cells = table.add_row().cells
            row_cells[0].text = f"Câu {question_id}"
            row_cells[1].text = q['question']
            
            # Format câu hỏi
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].paragraphs[0].runs[0].bold = True
            
            # Thêm các đáp án
            for letter in sorted(q['answers'].keys()):
                row_cells = table.add_row().cells
                row_cells[0].text = f"   {letter})"
                row_cells[1].text = q['answers'][letter]
            
            # Thêm đáp án đúng
            row_cells = table.add_row().cells
            row_cells[0].text = "   Đáp án"
            row_cells[1].text = q['correct_answer'] or 'Chưa xác định'
            
            # Format đáp án đúng
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].paragraphs[0].runs[0].bold = True
            
            # Không thêm hàng trống để phân cách - bỏ theo yêu cầu
        
        # Điều chỉnh độ rộng cột
        for i, width in enumerate([Inches(1.5), Inches(5)]):
            for row in table.rows:
                row.cells[i].width = width
        
        doc.save(file_path)
    
    def refresh_view(self):
        """Refresh toàn bộ view với error handling"""
        try:
            self.parse_questions()
            self.update_question_table()
            self.update_content_viewer()
            
            # Cập nhật status với thông tin chi tiết
            total_questions = len(self.questions)
            total_errors = len(getattr(self, 'parsing_errors', []))
            
            if total_errors > 0:
                status_msg = f"Đã refresh: {total_questions} câu hỏi hợp lệ, {total_errors} lỗi cần xử lý"
                self.update_status(status_msg)
                
                # Auto show errors nếu có
                if messagebox.askyesno("Có lỗi phát hiện", 
                                     f"Phát hiện {total_errors} câu hỏi có lỗi.\n"
                                     "Bạn có muốn xem và sửa ngay không?"):
                    self.show_parsing_errors()
            else:
                self.update_status(f"Đã refresh: {total_questions} câu hỏi hợp lệ")
                
        except Exception as e:
            error_msg = f"Lỗi khi refresh: {str(e)}"
            self.update_status(error_msg)
            messagebox.showerror("Lỗi", error_msg)
    
    def update_status(self, message):
        """Cập nhật status bar với thông tin chi tiết"""
        current_info = f" | {len(self.questions)} câu hợp lệ"
        if hasattr(self, 'parsing_errors') and self.parsing_errors:
            current_info += f", {len(self.parsing_errors)} lỗi"
        
        full_message = message + current_info
        self.status_label.config(text=full_message)
        print(f"📌 Status: {full_message}")  # Debug log
    
    def fix_question_numbers(self):
        """Sửa số thứ tự câu hỏi: Phát hiện câu dạng '226a' và chuyển thành số đúng"""
        if not self.raw_content:
            messagebox.showwarning("Cảnh báo", "Chưa có dữ liệu để xử lý!")
            return
        
        # Tìm các câu hỏi có chữ cái (ví dụ: 226a, 156b...)
        pattern_with_letter = re.compile(r'^Câu\s+(\d+)([a-zA-Z])[:\.]')
        fixes_needed = []
        
        for line_idx, line in enumerate(self.raw_content):
            match = pattern_with_letter.match(line.strip())
            if match:
                base_num = int(match.group(1))
                letter = match.group(2)
                fixes_needed.append({
                    'line_idx': line_idx,
                    'original_line': line,
                    'base_num': base_num,
                    'letter': letter,
                    'full_num': f"{base_num}{letter}"
                })
        
        if not fixes_needed:
            messagebox.showinfo("Thông báo", "Không tìm thấy câu hỏi nào có chữ cái cần sửa (ví dụ: 226a)")
            return
        
        # Hiển thị preview
        preview_text = f"Tìm thấy {len(fixes_needed)} câu hỏi cần sửa:\n\n"
        for fix in fixes_needed[:10]:  # Hiển thị tối đa 10 câu
            preview_text += f"• Dòng {fix['line_idx'] + 1}: Câu {fix['full_num']} → Câu {fix['base_num'] + 1}\n"
        
        if len(fixes_needed) > 10:
            preview_text += f"\n... và {len(fixes_needed) - 10} câu khác"
        
        preview_text += f"\n\nCác câu từ {fixes_needed[0]['base_num'] + 1} trở đi sẽ được tăng thêm 1."
        preview_text += f"\n\nBạn có muốn tiếp tục?"
        
        if not messagebox.askyesno("Xác nhận sửa số thứ tự", preview_text):
            return
        
        # Thực hiện sửa
        try:
            # Sắp xếp theo thứ tự ngược để sửa từ cuối lên đầu (tránh conflict)
            fixes_needed.sort(key=lambda x: x['base_num'], reverse=True)
            
            for fix in fixes_needed:
                base_num = fix['base_num']
                new_num = base_num + 1
                
                # Tăng tất cả câu từ new_num trở đi lên 1
                for i in range(len(self.raw_content)):
                    line = self.raw_content[i]
                    # Tìm câu hỏi >= new_num (không có chữ cái)
                    match = re.match(r'^Câu\s+(\d+)[:\.]', line.strip())
                    if match:
                        num = int(match.group(1))
                        if num >= new_num:
                            # Tăng số lên 1
                            self.raw_content[i] = re.sub(
                                r'^(Câu\s+)(\d+)([:\.])',
                                lambda m: f"{m.group(1)}{int(m.group(2)) + 1}{m.group(3)}",
                                line
                            )
                
                # Sửa câu có chữ cái: "226a" → "227"
                self.raw_content[fix['line_idx']] = re.sub(
                    r'^(Câu\s+)\d+[a-zA-Z]([:\.])',
                    f"\\g<1>{new_num}\\2",
                    fix['original_line']
                )
            
            # Parse lại
            self.parse_questions()
            self.update_question_table()
            self.update_content_viewer()
            
            messagebox.showinfo("Thành công", 
                              f"Đã sửa {len(fixes_needed)} câu hỏi!\n\n"
                              f"Vui lòng kiểm tra lại kết quả và lưu file nếu đúng.")
            self.update_status(f"Đã sửa {len(fixes_needed)} câu hỏi có chữ cái")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Có lỗi khi sửa số thứ tự:\n{str(e)}")

def main():
    root = tk.Tk()
    app = QuestionConverterGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()